"""
Document generation utilities for creating CSV, Word, and Excel documents.
Handles article exports and question document creation.
"""
import pandas as pd
from datetime import datetime
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

from storage import storage
from config import config


class DocumentGenerator:
    """Generates various document types for the pipeline."""
    
    def __init__(self):
        """Initialize document generator."""
        self.output_dir = config.BASE_DATA_PATH
    
    def generate_articles_csv(self, articles: List[Dict], filename: str = "articles.csv") -> str:
        """
        Generate CSV file with all article data.
        
        Args:
            articles: List of article dictionaries
            filename: Output filename
            
        Returns:
            Full path to generated CSV file
        """
        # Convert articles to DataFrame
        df = pd.DataFrame(articles)
        
        # Ensure required columns exist
        required_columns = ['title', 'url', 'source', 'license', 'date', 'summary', 'content']
        for col in required_columns:
            if col not in df.columns:
                df[col] = ''
        
        # Add metadata
        df['extracted_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Reorder columns
        column_order = ['title', 'url', 'source', 'license', 'date', 'summary', 'content', 'extracted_date']
        df = df[[col for col in column_order if col in df.columns]]
        
        # Save using storage abstraction
        file_path = storage.write_csv(df, filename)
        return file_path
    
    def generate_validated_csv(self, articles: List[Dict], filename: str = "validated_articles.csv") -> str:
        """
        Generate CSV file with validated articles including license information.
        
        Args:
            articles: List of validated article dictionaries
            filename: Output filename
            
        Returns:
            Full path to generated CSV file
        """
        df = pd.DataFrame(articles)
        
        # Save using storage abstraction
        file_path = storage.write_csv(df, filename)
        return file_path
    
    
    def _parse_habilidad(self, habilidad_str: str) -> Tuple[str, str]:
        """
        Parse habilidad string like '[Localizar-a]' or 'Interpretar-d' into components.
        
        Args:
            habilidad_str: String containing habilidad and tarea (e.g., '[Localizar-a]', 'Interpretar-d')
            
        Returns:
            Tuple of (habilidad, tarea_lectora)
        """
        # Remove brackets if present
        clean_str = habilidad_str.strip('[]').strip()
        
        # Split by hyphen
        if '-' in clean_str:
            parts = clean_str.split('-')
            habilidad = parts[0].strip()
            tarea = parts[1].strip() if len(parts) > 1 else ''
            return habilidad, tarea
        
        return clean_str, ''
    
    def generate_questions_word(self, article: Dict, questions: Dict, 
                               filename: str, is_improved: bool = False) -> str:
        """
        Generate Word document with questions in clean format (no numbers, no answers, no colors).
        
        Args:
            article: Article dictionary
            questions: Questions dictionary (parsed from agent response)
            filename: Output filename
            is_improved: Whether these are improved questions
            
        Returns:
            Full path to generated Word document
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Style configuration - Default to Times New Roman 14
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Title (16pt, Bold, Centered)
        title_para = doc.add_heading(article.get('title', 'Article Questions'), 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = None # Default color (black)
        
        # Add article text if available (from Agent 3 response)
        article_text = questions.get('article_text', '') if isinstance(questions, dict) else ''
        if article_text:
            # TEXTO DEL ARTÍCULO Heading (match Title style roughly or 14pt bold)
            text_heading = doc.add_paragraph('TEXTO DEL ARTÍCULO')
            text_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text_heading.runs[0].font.name = 'Times New Roman'
            text_heading.runs[0].font.size = Pt(14)
            text_heading.runs[0].bold = True
            
            doc.add_paragraph() # Spacer
                
            # Add the full text in multiple paragraphs
            paragraphs = article_text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Ensure font is applied explicitly
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                    
                    p.paragraph_format.space_after = Pt(12)
            
            # Source/Reference at the bottom (10pt as seen in analysis)
            if article.get('source') or article.get('author'):
                ref_text = f"{article.get('author', '')}. ({article.get('date', '')}). {article.get('title', '')}. {article.get('source', '')}."
                p_ref = doc.add_paragraph(ref_text)
                p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p_ref.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            
            doc.add_page_break()  # Start questions on new page
        else:
            print("[Document Generator] WARNING: No article text found in questions dict")
        
        # Questions
        if isinstance(questions, dict) and 'questions' in questions:
            questions_list = questions['questions']
        elif isinstance(questions, list):
            questions_list = questions
        else:
            questions_list = []
        
        for i, q in enumerate(questions_list, 1):
            if isinstance(q, dict):
                question_text = q.get('question', q.get('text', ''))
                alternatives = q.get('alternatives', {})
            else:
                question_text = str(q)
                alternatives = {}
            
            # Question text (14pt, Justified)
            p = doc.add_paragraph()
            run = p.add_run(question_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            
            # Alternatives (A-D) (14pt, Justified or Left)
            if alternatives:
                for letter in ['A', 'B', 'C', 'D']:
                    if letter in alternatives:
                        p = doc.add_paragraph()
                        # Hanging indent setup
                        p.paragraph_format.left_indent = Inches(0.3)
                        p.paragraph_format.first_line_indent = Inches(-0.3)
                        
                        run = p.add_run(f"{letter}) {alternatives[letter]}")
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_after = Pt(6)
            else:
                # If no alternatives dict, add warning
                p = doc.add_paragraph()
                p.add_run("[WARNING: No alternatives found for this question]")
            
            # Add page break after each question (one question per page)
            if i < len(questions_list):
                doc.add_page_break()
        
        # Save document
        full_path = os.path.join(self.output_dir, filename)
        os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else self.output_dir, exist_ok=True)
        doc.save(full_path)
        
        return full_path
    
    def generate_questions_excel(self, questions: Dict, filename: str) -> str:
        """
        Generate Excel file with question metadata and answers.
        
        Args:
            questions: Questions dictionary (parsed from agent response)
            filename: Output filename (should end with .xlsx)
            
        Returns:
            Full path to generated Excel file
        """
        # Prepare data for DataFrame
        excel_data = []
        
        # Questions
        if isinstance(questions, dict) and 'questions' in questions:
            questions_list = questions['questions']
        elif isinstance(questions, list):
            questions_list = questions
        else:
            questions_list = []
        
        for i, q in enumerate(questions_list, 1):
            if isinstance(q, dict):
                question_text = q.get('question', q.get('text', ''))
                clave = q.get('clave', q.get('answer', ''))
                justification = q.get('justification', q.get('explanation', ''))
                habilidad_raw = q.get('habilidad', '')
                
                # Parse habilidad into components
                habilidad, tarea_lectora = self._parse_habilidad(habilidad_raw)
            else:
                question_text = str(q)
                clave = ''
                justification = ''
                habilidad = ''
                tarea_lectora = ''
            
            # Add row to data
            excel_data.append({
                'Número de pregunta': i,
                'Clave': clave,
                'Habilidad': habilidad,
                'Tarea lectora': tarea_lectora,
                'Justificación': justification,
                'Acción': ''  # Always blank
            })
        
        # Create DataFrame
        df = pd.DataFrame(excel_data)
        
        # Save to Excel
        full_path = os.path.join(self.output_dir, filename)
        os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else self.output_dir, exist_ok=True)
        
        # Use storage abstraction for consistency
        df.to_excel(full_path, index=False, engine='openpyxl')
        
        print(f"[Document Generator] Excel file created: {full_path}")
        return full_path
    
    def parse_questions_from_response(self, response_text: str) -> Dict:
        """
        Parse questions from agent response text.
        
        Args:
            response_text: Raw text response from agent
            
        Returns:
            Dictionary with structured questions
        """
        questions = []
        lines = response_text.strip().split('\n')
        
        current_question = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('QUESTION'):
                # Save previous question if exists
                if current_question:
                    questions.append(current_question)
                    current_question = {}
                
                # Extract question text
                parts = line.split(':', 1)
                if len(parts) > 1:
                    current_question['question'] = parts[1].strip()
            
            elif line.startswith('ANSWER'):
                parts = line.split(':', 1)
                if len(parts) > 1:
                    current_question['answer'] = parts[1].strip()
            
            elif line.startswith('EXPLANATION'):
                parts = line.split(':', 1)
                if len(parts) > 1:
                    current_question['explanation'] = parts[1].strip()
        
        # Add last question
        if current_question:
            questions.append(current_question)
        
        return {'questions': questions}
    
    def merge_text_and_questions_docx(self, source_docx_path: str, 
                                      questions: Dict,
                                      output_path: str,
                                      title: str = "") -> str:
        """
        Merge source DOCX (article text) with generated questions into single DOCX.
        
        Final format:
        - Title (if provided)
        - Article text from source DOCX
        - Generated questions
        
        Args:
            source_docx_path: Path to source DOCX with article text
            questions: Questions dict from Agent 3
            output_path: Output path for merged DOCX
            title: Optional title to add at top
        
        Returns:
            Path to created file
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Load the source document directly to preserve formatting
        print(f"[DocGen] Loading source document: {source_docx_path}")
        print(f"[DocGen] Source exists: {os.path.exists(source_docx_path)}")
        
        merged_doc = Document(source_docx_path)
        print(f"[DocGen] Source loaded: {len(merged_doc.paragraphs)} paragraphs")
        
        # Add page break before questions
        merged_doc.add_page_break()
        
        # Add each question (1 per page, no numbers, no clave)
        questions_list = questions.get('questions', [])
        for i, q in enumerate(questions_list, 1):
            # Question text only (no number, no enumeration)
            q_para = merged_doc.add_paragraph(q.get('question', ''))
            q_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            q_para.paragraph_format.space_after = Pt(12)
            for run in q_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            
            # Alternatives - dict format {'A': 'text', 'B': 'text', ...}
            alternatives_dict = q.get('alternatives', {})
            for letter in ['A', 'B', 'C', 'D']:
                if letter in alternatives_dict:
                    alt_text = f"{letter}) {alternatives_dict[letter]}"
                    alt_para = merged_doc.add_paragraph(alt_text)
                    
                    # Hanging indent
                    alt_para.paragraph_format.left_indent = Inches(0.3)
                    alt_para.paragraph_format.first_line_indent = Inches(-0.3)
                    
                    alt_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    alt_para.paragraph_format.space_after = Pt(6)
                    for run in alt_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
            
            # Add page break after each question (except the last one)
            if i < len(questions_list):
                merged_doc.add_page_break()
        
        # Save merged document
        merged_doc.save(output_path)
        print(f"[DocGen] Merged DOCX saved: {output_path}")
        
        return output_path


# Global document generator instance
doc_generator = DocumentGenerator()

