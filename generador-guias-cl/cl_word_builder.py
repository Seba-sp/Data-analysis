"""
CL word assembly utilities.

This module parses a source docx into text section + question blocks using page breaks,
then assembles a final guide by keeping only selected questions per text.
"""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import BASE_DIR, CL_COLUMNS, CL_FILENAME_PREFIX


@dataclass
class ParsedTextDocument:
    """Represents one parsed text doc with body blocks split by pages."""

    text_blocks: List[object]
    question_blocks: List[List[object]]


def _element_has_page_break(element) -> bool:
    """Detect explicit page break in XML descendants."""
    for child in element.iter():
        if child.tag.endswith("br") and child.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
            return True
    return False


def _split_body_by_page_breaks(doc: Document) -> List[List[object]]:
    """Split body elements into page-like chunks using page breaks."""
    pages: List[List[object]] = []
    current: List[object] = []

    body_elements = list(doc.element.body)
    if body_elements and body_elements[-1].tag.endswith("sectPr"):
        body_elements = body_elements[:-1]

    for element in body_elements:
        current.append(deepcopy(element))
        if _element_has_page_break(element):
            pages.append(current)
            current = []

    if current:
        pages.append(current)

    return pages


def parse_cl_source_docx(docx_path: Path, expected_questions: int) -> ParsedTextDocument:
    """Parse CL source docx and split into text pages + question pages."""
    doc = Document(str(docx_path))
    pages = _split_body_by_page_breaks(doc)

    if len(pages) < expected_questions:
        raise ValueError(
            f"Documento {docx_path.name} tiene {len(pages)} bloques por salto de pagina, "
            f"pero se esperaban al menos {expected_questions}"
        )

    split_index = len(pages) - expected_questions
    if split_index < 1:
        split_index = 1

    text_pages = pages[:split_index]
    question_pages = pages[split_index:]

    text_blocks: List[object] = []
    for page in text_pages:
        text_blocks.extend(page)

    return ParsedTextDocument(text_blocks=text_blocks, question_blocks=question_pages)


def _append_blocks_to_document(target_doc: Document, blocks: List[object]) -> None:
    """Append XML blocks into target document body, before the sectPr."""
    body = target_doc.element.body
    sect_pr = body.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr")
    for block in blocks:
        copied = deepcopy(block)
        if sect_pr is not None:
            sect_pr.addprevious(copied)
        else:
            body.append(copied)


def _clean_default_paragraph(doc: Document) -> None:
    """Remove default empty paragraph from a new Document()."""
    body = doc.element.body
    if len(body) > 0 and body[0].tag.endswith("p"):
        p = body[0]
        if len(list(p.iter())) <= 2:
            body.remove(p)


def _blocks_end_with_page_break(blocks: List[object]) -> bool:
    """Return True if the last block carries an explicit page break."""
    if not blocks:
        return False
    return _element_has_page_break(blocks[-1])


def _insert_page_break_before_sectpr(doc: Document) -> None:
    """Insert a page-break paragraph before the sectPr element."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)

    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(p)
    else:
        body.append(p)


def _safe_filename_component(value: str) -> str:
    """Sanitize a label for Windows-safe filename use."""
    safe = "".join(ch for ch in str(value).strip() if ch not in '\\/:*?"<>|')
    safe = "_".join(safe.split())
    return safe or CL_FILENAME_PREFIX


def generate_cl_outputs(
    selected_texts: List[Dict[str, object]],
    by_codigo_df: Dict[str, pd.DataFrame],
    target_questions: int,
    guide_name: Optional[str] = None,
) -> Tuple[bytes, str, bytes, str, pd.DataFrame]:
    """
    Generate final CL Word + Excel report in memory.

    Returns:
        (word_bytes, word_filename, excel_bytes, excel_filename, report_df)
    """
    final_doc = Document()
    _clean_default_paragraph(final_doc)

    report_rows: List[Dict[str, object]] = []
    final_question_number = 1

    for text_idx, text_item in enumerate(selected_texts):
        codigo = str(text_item["Codigo Texto"])
        docx_path = BASE_DIR / Path(str(text_item["docx_path"]))
        removed_questions = set(text_item.get("removed_questions", []))

        question_df = by_codigo_df[codigo].sort_values(CL_COLUMNS["numero_pregunta"]).reset_index(drop=True)
        expected_questions = len(question_df)

        parsed = parse_cl_source_docx(docx_path, expected_questions=expected_questions)

        _append_blocks_to_document(final_doc, parsed.text_blocks)
        last_appended_blocks = parsed.text_blocks

        for idx, (_, row) in enumerate(question_df.iterrows(), start=1):
            original_num = int(row[CL_COLUMNS["numero_pregunta"]])
            if original_num in removed_questions:
                continue

            if idx - 1 >= len(parsed.question_blocks):
                continue

            kept_blocks = parsed.question_blocks[idx - 1]
            _append_blocks_to_document(final_doc, kept_blocks)
            last_appended_blocks = kept_blocks

            report_rows.append(
                {
                    "Pregunta final": final_question_number,
                    "Codigo Texto": codigo,
                    "N original del documento": original_num,
                    "Codigo Spot": row.get(CL_COLUMNS["codigo_spot"], ""),
                    "Clave": row.get(CL_COLUMNS["clave"], ""),
                    "Justificacion": row.get(CL_COLUMNS["justificacion"], ""),
                    "Habilidad": row.get(CL_COLUMNS["habilidad"], ""),
                    "Tarea lectora": row.get(CL_COLUMNS["tarea_lectora"], ""),
                }
            )
            final_question_number += 1

        # Ensure next text starts on a new page only when needed.
        if text_idx < len(selected_texts) - 1:
            if not _blocks_end_with_page_break(last_appended_blocks):
                _insert_page_break_before_sectpr(final_doc)

    report_df = pd.DataFrame(report_rows)

    if target_questions > 0 and len(report_df) != target_questions:
        raise ValueError(
            f"El total generado ({len(report_df)}) no coincide con el objetivo ({target_questions})."
        )

    file_prefix = _safe_filename_component(guide_name) if guide_name else CL_FILENAME_PREFIX
    word_filename = f"{file_prefix}.docx"
    excel_filename = f"{file_prefix}.xlsx"

    word_buffer = BytesIO()
    final_doc.save(word_buffer)
    word_bytes = word_buffer.getvalue()

    excel_buffer = BytesIO()
    report_df.to_excel(excel_buffer, index=False)
    excel_bytes = excel_buffer.getvalue()

    return word_bytes, word_filename, excel_bytes, excel_filename, report_df
