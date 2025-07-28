#!/usr/bin/env python3
"""
Report generator for individual student assessment reports
Generates PDF reports using Word templates
"""

import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx2pdf import convert
from storage import StorageClient

logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self):
        """Initialize report generator"""
        self.storage = StorageClient()
        self.template_path = "templates/plantilla_test_diagnostico.docx"
    
    def generate_individual_report(self, user: Dict[str, Any], assessment: Dict[str, Any], analysis_result: Dict[str, Any]) -> Path:
        """
        Generate individual PDF report for student
        
        Args:
            user: User information from webhook
            assessment: Assessment information from webhook
            analysis_result: Analysis results from AssessmentAnalyzer
            
        Returns:
            Path to generated PDF report
        """
        try:
            # Create temporary files for processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx, \
                 tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                
                temp_docx_path = temp_docx.name
                temp_pdf_path = temp_pdf.name
            
            # Load template
            template_doc = self._load_template()
            
            # Generate Word document
            self._generate_word_document(
                template_doc, 
                user, 
                assessment, 
                analysis_result, 
                temp_docx_path
            )
            
            # Convert to PDF
            convert(temp_docx_path, temp_pdf_path)
            
            # Save to storage
            report_filename = f"informe_{user.get('username', user['email'])}_{assessment.get('title', 'assessment')}.pdf"
            report_path = f"data/webhook_reports/{report_filename}"
            
            # Ensure directory exists
            self.storage.ensure_directory("data/webhook_reports/")
            
            # Read PDF and save to storage
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
                self.storage.write_bytes(report_path, pdf_content)
            
            # Clean up temporary files
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            
            logger.info(f"Generated report: {report_path}")
            return Path(report_path)
            
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}")
            raise
    
    def _load_template(self) -> Document:
        """Load Word template from storage"""
        try:
            if self.storage.exists(self.template_path):
                template_bytes = self.storage.read_bytes(self.template_path)
                
                # Create temporary file for docx library
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file.write(template_bytes)
                    temp_file_path = temp_file.name
                
                # Load document
                doc = Document(temp_file_path)
                
                # Clean up temporary file
                os.unlink(temp_file_path)
                
                return doc
            else:
                raise FileNotFoundError(f"Template not found: {self.template_path}")
                
        except Exception as e:
            logger.error(f"Error loading template: {str(e)}")
            raise
    
    def _generate_word_document(self, doc: Document, user: Dict[str, Any], assessment: Dict[str, Any], analysis_result: Dict[str, Any], output_path: str):
        """Generate Word document with student data"""
        try:
            # Extract data
            username = user.get("username", user["email"])
            level = analysis_result.get("level", "Nivel 2")
            percentage = analysis_result.get("overall_percentage", 0)
            area_results = analysis_result.get("area_results", {})
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(
                    paragraph, username, level, percentage, area_results
                )
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(
                                paragraph, username, level, percentage, area_results
                            )
            
            # Save document
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise
    
    def _replace_placeholders_in_paragraph(self, paragraph, username: str, level: str, percentage: float, area_results: Dict[str, Any]):
        """Replace placeholders in a paragraph"""
        text = paragraph.text
        
        # Replace basic placeholders
        text = text.replace("<<Nombre>>", username)
        text = text.replace("<<Nivel>>", level)
        text = text.replace("<<PD%>>", f"{percentage:.2f}%")
        
        # Replace area placeholders
        areas = list(area_results.keys())
        for i, area in enumerate(areas, 1):
            area_status = area_results[area]["status"]
            text = text.replace(f"<<T{i}>>", area_status)
        
        # Update paragraph text
        paragraph.text = text 