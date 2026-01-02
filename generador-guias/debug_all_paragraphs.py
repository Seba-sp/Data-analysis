from docx import Document
import sys

doc = Document(sys.argv[1])

with open('debug_all_output.txt', 'w', encoding='utf-8') as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n\n")
    f.write("ALL paragraphs (including empty):\n")
    f.write("=" * 100 + "\n")
    
    for idx in range(70, 95):
        if idx < len(doc.paragraphs):
            text = doc.paragraphs[idx].text
            if text.strip():
                f.write(f"[{idx:3d}] NONEMPTY: {text}\n")
            else:
                f.write(f"[{idx:3d}] EMPTY\n")

print("Output written to debug_all_output.txt")

