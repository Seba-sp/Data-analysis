"""
Word document processing module for splitting questions into individual files.
Each question is already 1 page, so we split by pages to preserve ALL formatting.
"""

from typing import List, Dict
import zipfile
import re
import tempfile
import os

from storage import StorageClient
from config import PREGUNTAS_DIVIDIDAS_DIR, SUBJECT_FOLDERS

class QuestionProcessor:
    """Handles splitting Word documents into individual question files by pages."""
    
    def __init__(self, storage_client: StorageClient):
        self.storage = storage_client
        
    def split_document_by_pages(self, docx_path: str) -> List[str]:
        """
        Split a Word document into individual pages (questions) by working directly with the ZIP structure.
        This preserves ALL formatting and images perfectly.
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            List of file paths to the created individual question files
        """
        try:
            # Read the Word document
            doc_bytes = self.storage.read_bytes(docx_path)
            
            # Extract the document content to find question boundaries
            from docx import Document
            from io import BytesIO
            original_doc = Document(BytesIO(doc_bytes))
            
            # Find question boundaries
            question_boundaries = self._find_question_boundaries(original_doc)
            
            if not question_boundaries:
                print("No question boundaries found")
                return []
            
            print(f"Found {len(question_boundaries)} questions")
            
            # Create individual files by copying the ZIP structure
            output_files = []
            for i, (start_idx, end_idx) in enumerate(question_boundaries):
                output_file = self._create_question_file_from_zip(
                    doc_bytes, start_idx, end_idx, i + 1
                )
                if output_file:
                    output_files.append(output_file)
            
            return output_files
            
        except Exception as e:
            print(f"Error splitting Word document {docx_path}: {e}")
            return []
    
    
    
    def _find_question_boundaries(self, doc) -> List[tuple]:
        """
        Find the boundaries of each question in the document.
        Always uses page-based splitting (1 question per page).
        
        Args:
            doc: Document object
            
        Returns:
            List of (start_index, end_index) tuples for each question
        """
        print("Using page-based splitting (1 question per page)")
        boundaries = self._find_page_based_boundaries(doc)
        
        return boundaries
    
    def _find_page_based_boundaries(self, doc) -> List[tuple]:
        """
        Find question boundaries based on page breaks.
        Since it's 1 question per page, split by page breaks.
        
        Args:
            doc: Document object
            
        Returns:
            List of (start_index, end_index) tuples for each question
        """
        boundaries = []
        current_start = 0
        page_breaks_found = 0
        
        for i, element in enumerate(doc.element.body):
            # Check if this element has a page break
            has_page_break = False
            
            # Look for page breaks in the element
            for child in element.iter():
                if child.tag.endswith('br'):
                    if child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                        has_page_break = True
                        page_breaks_found += 1
                        break
                elif child.tag.endswith('sectPr'):  # section properties often indicate page breaks
                    has_page_break = True
                    page_breaks_found += 1
                    break
            
            # If we found a page break, save current question and start new one
            if has_page_break:
                if i > current_start:  # Make sure we have content
                    boundaries.append((current_start, i - 1))
                current_start = i
        
        # Add the last question if it has content
        # Only add if we haven't already processed the last element
        if current_start < len(doc.element.body) - 1:
            boundaries.append((current_start, len(doc.element.body) - 1))
        elif current_start == len(doc.element.body) - 1:
            # The last element is just a section break, don't create a separate question
            pass
        
        # If no page breaks found, try to split by equal parts
        if not boundaries:
            print("No page breaks found, splitting by equal parts")
            total_elements = len(doc.element.body)
            if total_elements > 0:
                # Estimate number of questions (you might need to adjust this)
                estimated_questions = max(1, total_elements // 10)  # Assume ~10 elements per question
                elements_per_question = total_elements // estimated_questions
                
                for i in range(estimated_questions):
                    start_idx = i * elements_per_question
                    end_idx = min((i + 1) * elements_per_question - 1, total_elements - 1)
                    if start_idx <= end_idx:
                        boundaries.append((start_idx, end_idx))
        
        return boundaries
    
    def _create_question_file_from_zip(self, doc_bytes: bytes, start_idx: int, end_idx: int, question_num: int) -> str:
        """
        Create a question file using a hybrid approach.
        Creates a clean document structure while preserving images.
        
        Args:
            doc_bytes: Original document bytes
            start_idx: Start element index
            end_idx: End element index
            question_num: Question number
            
        Returns:
            Path to created file
        """
        try:
            # Create temporary directory
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract original ZIP
                original_zip_path = os.path.join(temp_dir, "original.docx")
                with open(original_zip_path, 'wb') as f:
                    f.write(doc_bytes)
                
                # Extract ZIP contents
                extract_dir = os.path.join(temp_dir, "extracted")
                with zipfile.ZipFile(original_zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                
                # Create a clean document.xml with only the question content
                self._create_clean_document_xml(extract_dir, start_idx, end_idx)
                
                # Create new ZIP file with all original files (preserving images)
                new_zip_path = os.path.join(temp_dir, f"question_{question_num}.docx")
                with zipfile.ZipFile(new_zip_path, 'w', zipfile.ZIP_DEFLATED) as new_zip:
                    # Add all files from the extracted directory
                    for root, dirs, files in os.walk(extract_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arc_name = os.path.relpath(file_path, extract_dir)
                            new_zip.write(file_path, arc_name)
                
                # Determine output path
                subject_folder = SUBJECT_FOLDERS.get("M1", "M1")  # Default to M1
                output_dir = PREGUNTAS_DIVIDIDAS_DIR / subject_folder
                self.storage.ensure_directory(str(output_dir))
                
                # Generate filename (we'll use a placeholder for now)
                filename = f"question_{question_num:03d}.docx"
                output_path = output_dir / filename
                
                # Copy the new file to output
                with open(new_zip_path, 'rb') as f:
                    self.storage.write_bytes(str(output_path), f.read())
                
                return str(output_path)
                
        except Exception as e:
            print(f"Error creating question file {question_num}: {e}")
            return ""
    
    def _create_clean_document_xml(self, extract_dir: str, start_idx: int, end_idx: int):
        """
        Create a clean document.xml with only the question content.
        This preserves images while creating a clean document structure.
        Also preserves A4 page size and adds proper margins (2.54 cm on all sides).
        
        Args:
            extract_dir: Directory where ZIP was extracted
            start_idx: Start element index
            end_idx: End element index
        """
        try:
            import xml.etree.ElementTree as ET
            
            # Read the original document.xml
            doc_xml_path = os.path.join(extract_dir, "word", "document.xml")
            tree = ET.parse(doc_xml_path)
            root = tree.getroot()
            
            # Find the body element
            body = None
            for elem in root.iter():
                if elem.tag.endswith('body'):
                    body = elem
                    break
            
            if body is None:
                print("Could not find body element")
                return
            
            # Get all direct child elements of body
            all_elements = list(body)
            
            # Keep only the elements in the specified range
            if start_idx < len(all_elements) and end_idx < len(all_elements):
                # Store the elements we want to keep
                elements_to_keep = []
                for i in range(start_idx, end_idx + 1):
                    if i < len(all_elements):
                        elements_to_keep.append(all_elements[i])
                
                # Clear the body completely
                body.clear()
                
                # Add back only the elements we want, cleaning them
                # Skip empty paragraphs at the beginning
                first_non_empty_found = False
                for element in elements_to_keep:
                    # Clean the element of page breaks and section properties
                    cleaned_element = self._clean_element_for_clean_document(element)
                    
                    # Check if this is an empty paragraph at the beginning
                    if not first_non_empty_found:
                        if self._is_empty_paragraph(cleaned_element):
                            continue  # Skip this empty paragraph
                        else:
                            first_non_empty_found = True
                    
                    body.append(cleaned_element)
                
                # Add proper section properties with A4 page size and margins
                self._add_section_properties(body)
            
            # Write back the modified XML
            tree.write(doc_xml_path, encoding='utf-8', xml_declaration=True)
            
        except Exception as e:
            print(f"Error creating clean document.xml: {e}")
    
    def _add_section_properties(self, body):
        """
        Add section properties with A4 page size and 2.54 cm margins on all sides.
        
        Args:
            body: The body element to add section properties to
        """
        try:
            import xml.etree.ElementTree as ET
            
            # Create section properties element
            sect_pr = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            
            # Add page size (A4: 21.0 cm x 29.7 cm)
            # Convert cm to twips (1 cm = 567 twips)
            page_width = int(21.0 * 567)  # 11907 twips
            page_height = int(29.7 * 567)  # 16840 twips
            
            pg_sz = ET.SubElement(sect_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgSz')
            pg_sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', str(page_width))
            pg_sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}h', str(page_height))
            
            # Add page margins (2.54 cm on all sides)
            # Convert cm to twips (1 cm = 567 twips)
            margin_twips = int(2.54 * 567)  # 1440 twips
            
            pg_mar = ET.SubElement(sect_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgMar')
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top', str(margin_twips))
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right', str(margin_twips))
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom', str(margin_twips))
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left', str(margin_twips))
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}header', '708')  # 1.25 cm
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footer', '708')  # 1.25 cm
            pg_mar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gutter', '0')
            
            # Add columns (single column)
            cols = ET.SubElement(sect_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols')
            cols.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '708')  # 1.25 cm
            
            # Add document grid
            doc_grid = ET.SubElement(sect_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docGrid')
            doc_grid.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}linePitch', '360')  # 6.35 mm
            
            # Append section properties to body
            body.append(sect_pr)
            
        except Exception as e:
            print(f"Error adding section properties: {e}")
    
    def _is_empty_paragraph(self, element):
        """
        Check if an element is an empty paragraph.
        
        Args:
            element: XML element to check
            
        Returns:
            True if the element is an empty paragraph, False otherwise
        """
        try:
            # Check if it's a paragraph element
            if not element.tag.endswith('p'):
                return False
            
            # Check if the paragraph has any text content
            for child in element.iter():
                if child.tag.endswith('t') and child.text and child.text.strip():
                    return False
            
            # Check if the paragraph has any non-text content (images, etc.)
            for child in element.iter():
                if child.tag.endswith('drawing') or child.tag.endswith('pict'):
                    return False
            
            return True
            
        except Exception as e:
            print(f"Warning: Could not check if paragraph is empty: {e}")
            return False
    
    def _clean_element_for_clean_document(self, element):
        """
        Clean an element to ensure it works properly in a clean document.
        Removes page breaks, section properties, and other elements that might cause issues.
        
        Args:
            element: XML element to clean
            
        Returns:
            Cleaned element (always returns the element, never None)
        """
        try:
            # Find and remove problematic elements
            children_to_remove = []
            for child in list(element):
                # Remove page breaks
                if child.tag.endswith('br'):
                    if child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                        children_to_remove.append(child)
                # Remove section properties
                elif child.tag.endswith('sectPr'):
                    children_to_remove.append(child)
                # Recursively clean children
                else:
                    self._clean_element_for_clean_document(child)
            
            # Remove the identified problematic elements
            for child in children_to_remove:
                element.remove(child)
            
            # Always return the element
            return element
                    
        except Exception as e:
            print(f"Warning: Could not clean element: {e}")
            # Always return the original element even if cleaning failed
            return element

    def process_word_document(
        self, 
        docx_path: str, 
        excel_data: List[Dict], 
        subject: str
    ) -> List[Dict[str, str]]:
        """
        Process a complete Word document and create individual question files.
        
        Args:
            docx_path: Path to the Word document
            excel_data: List of dictionaries with question metadata
            subject: Subject area
            
        Returns:
            List of dictionaries with processing results
        """
        # Split document by pages using the new ZIP-based approach
        question_files = self.split_document_by_pages(docx_path)
        
        if len(question_files) != len(excel_data):
            print(f"Warning: Number of questions ({len(question_files)}) doesn't match Excel rows ({len(excel_data)})")
        
        results = []
        
        for i, (question_file, excel_row) in enumerate(zip(question_files, excel_data)):
            try:
                # Get pregunta_id from excel_row (should be generated beforehand)
                pregunta_id = excel_row.get('PreguntaID', f'Q{i+1:03d}')
                
                # Rename the file with the proper PreguntaID
                if question_file:
                    # Determine output directory
                    subject_folder = SUBJECT_FOLDERS.get(subject, subject.upper())
                    output_dir = PREGUNTAS_DIVIDIDAS_DIR / subject_folder
                    
                    # Create new filename with PreguntaID
                    new_filename = f"{pregunta_id}.docx"
                    new_file_path = output_dir / new_filename
                    
                    # Copy the file with the new name
                    with open(question_file, 'rb') as f:
                        self.storage.write_bytes(str(new_file_path), f.read())
                    
                    # Delete the temporary file
                    os.remove(question_file)
                    
                    results.append({
                        'pregunta_id': pregunta_id,
                        'question_number': i + 1,
                        'file_path': str(new_file_path),
                        'filename': new_filename,
                        'success': True
                    })
                else:
                    results.append({
                        'pregunta_id': pregunta_id,
                        'question_number': i + 1,
                        'file_path': '',
                        'filename': '',
                        'success': False,
                        'error': 'Failed to create file'
                    })
                    
            except Exception as e:
                results.append({
                    'pregunta_id': excel_row.get('PreguntaID', f'Q{i+1:03d}'),
                    'question_number': i + 1,
                    'file_path': '',
                    'filename': '',
                    'success': False,
                    'error': str(e)
                })
        
        return results

# Test the processor
if __name__ == "__main__":
    from storage import StorageClient
    
    storage = StorageClient()
    processor = QuestionProcessor(storage)
    
    # Test with sample file if it exists
    sample_file = "sets/Ensayo Agosto 2025 - Física.docx"
    if storage.exists(sample_file):
        page_docs = processor.split_document_by_pages(sample_file)
        print(f"Split {sample_file} into {len(page_docs)} pages")
        for i, doc in enumerate(page_docs[:3]):  # Show first 3
            # Get first paragraph text as preview
            preview = ""
            if doc.paragraphs:
                preview = doc.paragraphs[0].text[:100]
            print(f"Page {i+1}: {preview}...")
    else:
        print(f"Sample file {sample_file} not found")
