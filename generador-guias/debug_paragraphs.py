from docx import Document
import sys

doc = Document(sys.argv[1])

with open('debug_output.txt', 'w', encoding='utf-8') as f:
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n\n")
    f.write("All non-empty paragraphs with indices:\n")
    f.write("=" * 100 + "\n")
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            f.write(f"[{idx:3d}] {text}\n")

print("Output written to debug_output.txt")

