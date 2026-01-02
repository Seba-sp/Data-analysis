import re
import unicodedata
import argparse
from docx import Document
from pathlib import Path

def normalize_text(text):
    # Remove accents, lowercase, strip, and collapse spaces
    text = re.sub(r'^Pregunta\s*\d+\s*', '', text, flags=re.IGNORECASE)
    text = text.lower().strip()
    text = ''.join(c for c in unicodedata.normalize('NFD', text)
                  if unicodedata.category(c) != 'Mn')
    text = re.sub(r'\s+', ' ', text)
    return text

def get_page_number(doc, paragraph_index):
    """
    Try to determine page number using multiple methods:
    1. Look for explicit 'Pregunta N' text near the question
    2. Count actual page breaks in the document
    3. Fall back to estimation if no explicit indicators
    """
    # Method 1: Look for "Pregunta N" pattern near this paragraph
    search_start = max(0, paragraph_index - 5)
    search_end = min(len(doc.paragraphs), paragraph_index + 2)
    
    for i in range(search_start, search_end):
        text = doc.paragraphs[i].text.strip()
        # Match patterns like "Pregunta 1", "Pregunta 2.", "1.", "2)", etc.
        match = re.match(r'^(?:Pregunta\s*)?(\d+)[.:\)]?\s*$', text, re.IGNORECASE)
        if match:
            try:
                question_num = int(match.group(1))
                # Assume one question per page (common format)
                return question_num
            except:
                pass
    
    # Method 2: Count actual page breaks
    page_num = 1
    page_breaks_found = False
    
    for i in range(min(paragraph_index, len(doc.paragraphs))):
        paragraph = doc.paragraphs[i]
        
        # Check paragraph XML for page breaks
        try:
            xml_str = str(paragraph._element.xml)
            # Look for page break elements
            if 'w:type="page"' in xml_str:
                page_num += 1
                page_breaks_found = True
        except:
            pass
    
    # If we found page breaks, use that count
    if page_breaks_found:
        return page_num
    
    # Method 3: Fallback - estimate based on paragraph density
    # Use a more conservative estimate (40 paragraphs per page)
    # This accounts for questions with lots of spacing
    return (paragraph_index // 40) + 1

def extract_questions_from_tables(doc, debug=False):
    """Extract questions that are formatted in tables"""
    questions = []
    
    # Build a map of table positions relative to paragraphs
    # This helps us estimate which "paragraph index" a table would be at
    table_positions = {}
    para_idx = 0
    
    for element in doc.element.body:
        if element.tag.endswith('tbl'):  # It's a table
            table_num = len(table_positions)
            table_positions[table_num] = para_idx
        elif element.tag.endswith('p'):  # It's a paragraph
            para_idx += 1
    
    for table_idx, table in enumerate(doc.tables):
        # Look for tables with A), B), C), D) or A), B), C), D), E) in first column
        if table.rows and len(table.rows) >= 4:
            first_col_texts = []
            try:
                for row in table.rows:
                    if row.cells:
                        first_col_texts.append(row.cells[0].text.strip())
            except:
                continue
            
            # Check if this looks like a question table (has A) and D) or E))
            has_a = any(re.match(r'^[aA][\s\u00A0]*\)', text) for text in first_col_texts)
            has_d = any(re.match(r'^[dD][\s\u00A0]*\)', text) for text in first_col_texts)
            has_e = any(re.match(r'^[eE][\s\u00A0]*\)', text) for text in first_col_texts)
            
            if has_a and (has_d or has_e):
                # Extract alternatives from table
                alternatives = []
                for row in table.rows:
                    if row.cells:
                        first_cell = row.cells[0].text.strip()
                        if re.match(r'^[aA][\s\u00A0]*\)', first_cell):
                            # Collect all cell texts in this row
                            row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            alternatives.append(row_text)
                        elif re.match(r'^[bB][\s\u00A0]*\)', first_cell):
                            row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            alternatives.append(row_text)
                        elif re.match(r'^[cC][\s\u00A0]*\)', first_cell):
                            row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            alternatives.append(row_text)
                        elif re.match(r'^[dD][\s\u00A0]*\)', first_cell):
                            row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            alternatives.append(row_text)
                        elif re.match(r'^[eE][\s\u00A0]*\)', first_cell):
                            row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            alternatives.append(row_text)
                
                if len(alternatives) >= 4:  # Accept 4 or 5 alternatives
                    # Try to find question text before the table
                    enunciado = f"Pregunta en tabla {table_idx + 1}"  # Placeholder
                    
                    # Calculate page number based on table position
                    para_position = table_positions.get(table_idx, 0)
                    page_num = get_page_number(doc, para_position)
                    
                    questions.append((enunciado, alternatives, -1, page_num))
                    
                    if debug:
                        print(f"\n[DEBUG] Found question in table {table_idx + 1} with {len(alternatives)} alternatives")
                        for i, alt in enumerate(alternatives):
                            safe_alt = alt[:60].encode('utf-8', errors='replace').decode('utf-8')
                            print(f"  [DEBUG] Alt {chr(65+i)}: '{safe_alt}'")
    
    return questions

def extract_questions_with_alternatives(docx_path, debug=False):
    doc = Document(docx_path)
    questions = []
    i = 0
    
    if debug:
        try:
            print(f"\n[DEBUG] Total paragraphs in document: {len(doc.paragraphs)}")
            print(f"[DEBUG] Total tables in document: {len(doc.tables)}")
            print("[DEBUG] Detecting questions with 4 or 5 alternatives (A-D or A-E)")
            print("[DEBUG] First 30 non-empty paragraphs:")
            count = 0
            for idx, p in enumerate(doc.paragraphs):
                if p.text.strip() and count < 30:
                    # Handle encoding issues for console output
                    safe_text = p.text.strip()[:100].encode('utf-8', errors='replace').decode('utf-8')
                    print(f"  [{idx}] '{safe_text}'")
                    count += 1
        except Exception as e:
            print(f"[DEBUG] Error printing paragraphs: {e}")
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        # Check if this paragraph starts an alternative (A, B, C, or D)
        alt_match = re.match(r'^[aA][\s\u00A0]*\)', text)
        
        if alt_match and i > 0:
            # Found alternative A - backtrack to find the question
            # The question is all non-empty paragraphs before this that aren't alternatives
            question_start_idx = i - 1
            enunciado_lines = []
            
            # Go backwards to collect question text
            back_idx = i - 1
            empty_count = 0
            while back_idx >= 0:
                back_text = doc.paragraphs[back_idx].text.strip()
                # Stop if we hit another alternative set (D or E)
                if re.match(r'^[dDeE][\s\u00A0]*\)', back_text):
                    break
                if back_text:
                    # Skip source citations (like "Recuperado:", "Fuente:", etc.)
                    if re.match(r'^(Recuperado|Fuente|Referencia|Bibliograf[ií]a|Imagen|Tabla|Figura)[\s:]*', back_text, re.IGNORECASE):
                        back_idx -= 1
                        empty_count = 0  # Reset empty counter when we skip citations
                        continue
                    enunciado_lines.insert(0, back_text)
                    question_start_idx = back_idx
                    empty_count = 0  # Reset empty counter when we find text
                else:
                    empty_count += 1
                    # Stop if we hit 5+ consecutive empty paragraphs (indicates question boundary/page break)
                    # Increased to 5 to better support one-question-per-page format
                    if empty_count >= 5:
                        break
                back_idx -= 1
                # Increased limit for one-question-per-page format
                if len(enunciado_lines) > 20:
                    break
            
            enunciado = ' '.join(enunciado_lines)
            
            if debug:
                safe_text = enunciado[:80].encode('utf-8', errors='replace').decode('utf-8')
                print(f"\n[DEBUG] Found question at paragraph {question_start_idx}: '{safe_text}'")
            
            # Now collect the alternatives (4 or 5) starting from current position
            alternatives = []
            j = i
            alt_patterns = [
                (r'^[aA][\s\u00A0]*\)', 'A'),
                (r'^[bB][\s\u00A0]*\)', 'B'),
                (r'^[cC][\s\u00A0]*\)', 'C'),
                (r'^[dD][\s\u00A0]*\)', 'D'),
                (r'^[eE][\s\u00A0]*\)', 'E')
            ]
            
            expected_alt_idx = 0
            empty_between_alts = 0
            max_search_distance = 25  # Maximum paragraphs to search for all alternatives
            
            while j < len(doc.paragraphs) and expected_alt_idx < 5 and (j - i) < max_search_distance:
                alt_text = doc.paragraphs[j].text.strip()
                pattern, letter = alt_patterns[expected_alt_idx]
                
                if re.match(pattern, alt_text):
                    alternatives.append(alt_text)
                    if debug:
                        safe_alt = alt_text[:60].encode('utf-8', errors='replace').decode('utf-8')
                        print(f"  [DEBUG] Found alternative {letter}: '{safe_alt}'")
                    expected_alt_idx += 1
                    empty_between_alts = 0  # Reset counter when we find an alternative
                    j += 1
                elif not alt_text:
                    # Skip empty paragraphs between alternatives
                    empty_between_alts += 1
                    # Allow up to 3 empty paragraphs between alternatives (for one-question-per-page)
                    if empty_between_alts > 3:
                        if debug:
                            print(f"  [DEBUG] Too many empty paragraphs ({empty_between_alts}), stopping search")
                        break
                    j += 1
                else:
                    # If we expected an alternative but found something else, stop
                    break
            
            # Accept questions with 4 or 5 alternatives
            if len(alternatives) >= 4 and enunciado:
                page_num = get_page_number(doc, question_start_idx)
                questions.append((enunciado, alternatives, question_start_idx, page_num))
                if debug:
                    print(f"  [DEBUG] [OK] Question saved with {len(alternatives)} alternatives")
                i = j
            else:
                if debug:
                    print(f"  [DEBUG] [SKIP] Question skipped (found {len(alternatives)} alternatives, enunciado length: {len(enunciado)})")
                i += 1
        else:
            i += 1
    
    # Also extract questions from tables
    table_questions = extract_questions_from_tables(doc, debug)
    questions.extend(table_questions)
    
    return questions

def alternatives_equal(alts1, alts2):
    # Both must have same number of alternatives (4 or 5)
    if len(alts1) != len(alts2):
        return False
    # Must have at least 4 alternatives
    if len(alts1) < 4:
        return False
    for a, b in zip(alts1, alts2):
        if normalize_text(a) != normalize_text(b):
            return False
    return True

def levenshtein(s1, s2):
    # Simple Levenshtein distance for possible matches
    if len(s1) < len(s2):
        return levenshtein(s2, s1)
    if len(s2) == 0:
        return len(s1)
    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    return previous_row[-1]

def compare_questions(file1_questions, file2_questions, file1_name, file2_name, similarity_threshold):
    """Compare questions between two files and return matches"""
    repeated_questions = []
    possible_matches = []
    
    for tenunciado, talts, tidx, tpage in file1_questions:
        norm_tenunciado = normalize_text(tenunciado)
        found = False
        
        for penunciado, palts, pidx, ppage in file2_questions:
            norm_penunciado = normalize_text(penunciado)
            
            # Exact match
            if norm_tenunciado == norm_penunciado and alternatives_equal(talts, palts):
                repeated_questions.append({
                    'file1_name': file1_name,
                    'file2_name': file2_name,
                    'file1_page': tpage,
                    'file2_page': ppage,
                    'question_text': penunciado,
                    'alternatives': palts
                })
                found = True
                break
        
        if not found:
            # Check for possible match (enunciado very similar, alternatives equal)
            for penunciado, palts, pidx, ppage in file2_questions:
                norm_penunciado = normalize_text(penunciado)
                lev = levenshtein(norm_tenunciado, norm_penunciado)
                max_len = max(len(norm_tenunciado), len(norm_penunciado))
                
                if max_len > 0 and lev <= max(2, int(similarity_threshold * max_len)) and alternatives_equal(talts, palts):
                    possible_matches.append({
                        'file1_name': file1_name,
                        'file2_name': file2_name,
                        'file1_page': tpage,
                        'file2_page': ppage,
                        'file1_question_text': tenunciado,
                        'file2_question_text': penunciado,
                        'alternatives': palts,
                        'levenshtein': lev
                    })
                    break
    
    return repeated_questions, possible_matches

def generate_report(repeated_questions, possible_matches, output_path):
    """Generate Word document report with duplicates"""
    doc = Document()
    doc.add_heading('Duplicate Questions Report', 0)
    
    if repeated_questions:
        doc.add_heading('Exact Duplicates', level=1)
        for rq in repeated_questions:
            doc.add_paragraph(f"File 1: {rq['file1_name']} (Page {rq['file1_page']})")
            doc.add_paragraph(f"File 2: {rq['file2_name']} (Page {rq['file2_page']})")
            doc.add_paragraph(f"Question: {rq['question_text']}")
            for alt in rq['alternatives']:
                doc.add_paragraph(alt, style='List Bullet')
            doc.add_paragraph('---')
    else:
        doc.add_paragraph('No exact duplicate questions found.')
    
    if possible_matches:
        doc.add_heading('Possible Matches (Review Manually)', level=1)
        for pm in possible_matches:
            doc.add_paragraph(f"File 1: {pm['file1_name']} (Page {pm['file1_page']})")
            doc.add_paragraph(f"File 2: {pm['file2_name']} (Page {pm['file2_page']})")
            doc.add_paragraph('Question in File 1:')
            doc.add_paragraph(pm['file1_question_text'])
            doc.add_paragraph('Question in File 2:')
            doc.add_paragraph(pm['file2_question_text'])
            doc.add_paragraph(f"Levenshtein Distance: {pm['levenshtein']}")
            for alt in pm['alternatives']:
                doc.add_paragraph(alt, style='List Bullet')
            doc.add_paragraph('---')
    
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")

def generate_text_summary(repeated_questions, possible_matches, output_path):
    """Generate text summary"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('=== DUPLICATE QUESTIONS SUMMARY ===\n\n')
        
        f.write(f'Total exact duplicates found: {len(repeated_questions)}\n')
        f.write(f'Total possible matches found: {len(possible_matches)}\n\n')
        
        if repeated_questions:
            f.write('EXACT DUPLICATES:\n')
            f.write('-' * 80 + '\n')
            for rq in repeated_questions:
                f.write(f"\n✓ Question: '{rq['question_text'][:100]}...'\n")
                f.write(f"  File 1: {rq['file1_name']} (Page {rq['file1_page']})\n")
                f.write(f"  File 2: {rq['file2_name']} (Page {rq['file2_page']})\n")
        
        if possible_matches:
            f.write('\n\nPOSSIBLE MATCHES (Review Manually):\n')
            f.write('-' * 80 + '\n')
            for pm in possible_matches:
                f.write(f"\n~ File 1 Question: '{pm['file1_question_text'][:100]}...'\n")
                f.write(f"  File 2 Question: '{pm['file2_question_text'][:100]}...'\n")
                f.write(f"  File 1: {pm['file1_name']} (Page {pm['file1_page']})\n")
                f.write(f"  File 2: {pm['file2_name']} (Page {pm['file2_page']})\n")
                f.write(f"  Levenshtein Distance: {pm['levenshtein']}\n")
    
    print(f"Summary saved to: {output_path}")

def main():
    # Configuration parameters
    SIMILARITY_THRESHOLD = 0.04  # 4% difference allowed for possible matches
    
    parser = argparse.ArgumentParser(
        description='Compare Word documents to find duplicate questions (supports 4 or 5 alternatives: A-D or A-E)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Compare two specific files:
  python detectar_duplicados.py --file1 "document1.docx" --file2 "document2.docx"
  
  # Compare one file with all files in a folder:
  python detectar_duplicados.py --file1 "document1.docx" --folder "path/to/folder"
  
  # With debug mode to see what questions are detected:
  python detectar_duplicados.py --file1 "document1.docx" --folder "path/to/folder" --debug
  
Note: This tool detects questions with 4 alternatives (A-D) or 5 alternatives (A-E).
      It also handles one-question-per-page formats with more spacing.
        """
    )
    
    parser.add_argument('--file1', required=True, help='Path to first file to compare')
    parser.add_argument('--file2', help='Path to second file to compare (use this OR --folder)')
    parser.add_argument('--folder', help='Path to folder containing files to compare against file1 (use this OR --file2)')
    parser.add_argument('--output', default='duplicate_report', help='Output file name (without extension, default: duplicate_report)')
    parser.add_argument('--threshold', type=float, default=SIMILARITY_THRESHOLD, help=f'Similarity threshold for possible matches (default: {SIMILARITY_THRESHOLD})')
    parser.add_argument('--debug', action='store_true', help='Enable debug output to see what questions are detected')
    
    args = parser.parse_args()
    
    # Validate arguments
    if not args.file2 and not args.folder:
        parser.error("You must specify either --file2 or --folder")
    
    if args.file2 and args.folder:
        parser.error("Specify only --file2 OR --folder, not both")
    
    # Check if file1 exists
    file1_path = Path(args.file1)
    if not file1_path.exists():
        print(f"Error: File not found: {args.file1}")
        return
    
    print(f"\nLoading questions from: {file1_path.name}")
    file1_questions = extract_questions_with_alternatives(str(file1_path), debug=args.debug)
    print(f"  Found {len(file1_questions)} questions")
    
    all_repeated = []
    all_possible = []
    
    # Mode 1: File to File comparison
    if args.file2:
        file2_path = Path(args.file2)
        if not file2_path.exists():
            print(f"Error: File not found: {args.file2}")
            return
        
        print(f"\nComparing with: {file2_path.name}")
        file2_questions = extract_questions_with_alternatives(str(file2_path), debug=args.debug)
        print(f"  Found {len(file2_questions)} questions")
        
        repeated, possible = compare_questions(
            file1_questions, file2_questions,
            file1_path.name, file2_path.name,
            args.threshold
        )
        all_repeated.extend(repeated)
        all_possible.extend(possible)
    
    # Mode 2: File to Folder comparison
    else:
        folder_path = Path(args.folder)
        if not folder_path.exists() or not folder_path.is_dir():
            print(f"Error: Folder not found: {args.folder}")
            return
        
        docx_files = list(folder_path.glob('*.docx'))
        # Exclude temporary Word files that start with ~$
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        
        if not docx_files:
            print(f"No .docx files found in: {args.folder}")
            return
        
        print(f"\nFound {len(docx_files)} files in folder to compare")
        
        for file2_path in docx_files:
            # Skip comparing file with itself
            if file2_path.resolve() == file1_path.resolve():
                continue
            
            print(f"\nComparing with: {file2_path.name}")
            file2_questions = extract_questions_with_alternatives(str(file2_path), debug=args.debug)
            print(f"  Found {len(file2_questions)} questions")
            
            repeated, possible = compare_questions(
                file1_questions, file2_questions,
                file1_path.name, file2_path.name,
                args.threshold
            )
            all_repeated.extend(repeated)
            all_possible.extend(possible)
    
    # Generate reports
    print(f"\n{'=' * 60}")
    print(f"RESULTS:")
    print(f"  Exact duplicates: {len(all_repeated)}")
    print(f"  Possible matches: {len(all_possible)}")
    print(f"{'=' * 60}")
    
    output_dir = file1_path.parent
    report_docx = output_dir / f"{args.output}.docx"
    report_txt = output_dir / f"{args.output}.txt"
    
    generate_report(all_repeated, all_possible, str(report_docx))
    generate_text_summary(all_repeated, all_possible, str(report_txt))
    
    print("\n[OK] Processing complete!")

if __name__ == "__main__":
    main()
